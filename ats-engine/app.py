from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
import spacy
import base64
import tempfile

import pdfplumber
import docx
import os
import PyPDF2
import io
from PIL import Image
import pytesseract
import json
from datetime import datetime
import logging
from werkzeug.utils import secure_filename

from scorer.scorer import score_resume
from utils.pdf_reader import extract_text_from_pdf
from utils.docx_reader import extract_text_from_docx
from scorer.keyword_extractor import extract_keywords

app = Flask(__name__)
CORS(app)  # Enable modern CORS support
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = 'temp_uploads'

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

nlp = spacy.load("en_core_web_sm")

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def extract_text_from_pdf(file):
    """Extract text from PDF using multiple modern methods with fallbacks"""
    file.seek(0)
    text = ""
    extraction_method = ""
    
    logger.info(f"Starting PDF extraction for file: {file.filename}")
    
    # Method 1: Try pdfplumber first (best for standard PDFs)
    try:
        with pdfplumber.open(file) as pdf:
            pages_text = []
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    pages_text.append(page_text)
                logger.info(f"Page {i+1}: {len(page_text) if page_text else 0} characters")
            text = "\n".join(pages_text)
            if text.strip():
                extraction_method = "pdfplumber"
                logger.info(f"✓ pdfplumber extraction successful: {len(text)} chars")
                return text, extraction_method
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")
    
    # Method 2: Try PyPDF2 as backup
    try:
        file.seek(0)
        reader = PyPDF2.PdfReader(file)
        pages_text = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                pages_text.append(page_text)
        text = "\n".join(pages_text)
        if text.strip():
            extraction_method = "PyPDF2"
            logger.info(f"✓ PyPDF2 extraction successful: {len(text)} chars")
            return text, extraction_method
    except Exception as e:
        logger.warning(f"PyPDF2 failed: {e}")
    
    # Method 3: OCR fallback for image-based or stylized PDFs
    try:
        file.seek(0)
        import fitz  # PyMuPDF
        pdf_document = fitz.open(stream=file.read(), filetype="pdf")
        ocr_text = []
        
        logger.info(f"Attempting OCR extraction for {len(pdf_document)} pages")
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            # Convert page to image with higher DPI for better OCR
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality
            img_data = pix.tobytes("png")
            image = Image.open(io.BytesIO(img_data))
            
            # Use OCR to extract text with better configuration
            custom_config = r'--oem 3 --psm 6'
            page_text = pytesseract.image_to_string(image, config=custom_config)
            if page_text.strip():
                ocr_text.append(page_text)
                logger.info(f"OCR Page {page_num+1}: {len(page_text)} characters")
        
        text = "\n".join(ocr_text)
        if text.strip():
            extraction_method = "OCR"
            logger.info(f"✓ OCR extraction successful: {len(text)} chars")
            return text, extraction_method
            
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
    
    return text, "failed"

def extract_skills(text):
    doc = nlp(text.lower())
    return {token.text for token in doc if token.pos_ in ["NOUN", "PROPN"]}

@app.route('/health')
def health_check():
    """Health check endpoint for monitoring"""
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "2.0.0",
        "features": {
            "pdf_extraction": True,
            "ocr_support": True,
            "docx_support": True,
            "ai_keyword_extraction": True
        }
    })

@app.route('/api/info')
def api_info():
    """API information endpoint"""
    return jsonify({
        "name": "SmartCV ATS Engine",
        "version": "2.0.0",
        "description": "Next-generation AI-powered resume analysis and ATS scoring",
        "endpoints": {
            "/": "Main application interface",
            "/score": "Resume scoring and analysis",
            "/extract-text": "Text extraction from documents",
            "/health": "Health check",
            "/api/info": "API information"
        },
        "supported_formats": ["PDF", "DOCX", "Text"],
        "extraction_methods": ["pdfplumber", "PyPDF2", "OCR"],
        "max_file_size": "16MB"
    })

@app.route("/")
def home():
    return render_template("score_resume.html")

@app.route('/extract-text', methods=['POST'])
def extract_text():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    text = ""

    if file.filename.endswith(".pdf"):
        with pdfplumber.open(file) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)

    elif file.filename.endswith(".docx"):
        doc = docx.Document(file)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    else:
        return jsonify({"error": "Unsupported file type"}), 400

    return jsonify({"text": text})

@app.route('/score', methods=['POST'])
def score():
    start_time = datetime.now()
    extraction_method = ""
    
    try:
        # Support both text and file uploads
        resume_text = request.form.get("resume_text")
        logger.info(f"Text area content length: {len(resume_text) if resume_text else 0}")
        logger.info(f"Files in request: {list(request.files.keys())}")
        
        if not resume_text and "resume_file" in request.files:
            file = request.files["resume_file"]
            logger.info(f"File: {file.filename}, Content-Type: {getattr(file, 'content_type', 'N/A')}")
            
            if file and file.filename and file.filename != '':
                # Secure filename
                filename = secure_filename(file.filename)
                logger.info(f"Processing file: {filename}")
                
                if file.filename.lower().endswith(".pdf"):
                    try:
                        resume_text, extraction_method = extract_text_from_pdf(file)
                        logger.info(f"Extraction method: {extraction_method}, Text length: {len(resume_text) if resume_text else 0}")
                        
                        if resume_text:
                            logger.info(f"Preview: {resume_text[:200]}...")
                    except Exception as pdf_error:
                        logger.error(f"PDF extraction error: {pdf_error}")
                        return jsonify({
                            "error": f"Failed to extract text from PDF: {str(pdf_error)}",
                            "extraction_method": "failed"
                        }), 400
                        
                elif file.filename.lower().endswith(".docx"):
                    try:
                        file.seek(0)
                        doc = docx.Document(file)
                        resume_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                        extraction_method = "docx"
                        logger.info(f"DOCX extracted text length: {len(resume_text) if resume_text else 0}")
                    except Exception as docx_error:
                        logger.error(f"DOCX extraction error: {docx_error}")
                        return jsonify({
                            "error": f"Failed to extract text from DOCX: {str(docx_error)}",
                            "extraction_method": "failed"
                        }), 400
                else:
                    return jsonify({"error": "Unsupported file type. Please upload PDF or DOCX files."}), 400
            else:
                logger.warning("No valid file found")
                return jsonify({"error": "No file selected or file is empty"}), 400

        logger.info(f"Final text check - Length: {len(resume_text) if resume_text else 0}")
        
        if not resume_text or not resume_text.strip():
            return jsonify({"error": "No resume text provided or file appears to be empty"}), 400

        # Enhanced keyword extraction with job-specific intelligence
        job_keywords = request.form.get("keywords")
        if not job_keywords:
            # Enhanced keyword extraction using NLP
            doc = nlp(resume_text.lower())
            
            # Extract skills, technologies, and important nouns
            keywords_set = set()
            
            # Add technical skills and proper nouns
            for token in doc:
                if (token.pos_ in ["NOUN", "PROPN"] and 
                    len(token.text) > 2 and 
                    not token.is_stop and 
                    token.is_alpha):
                    keywords_set.add(token.text)
            
            # Add named entities (technologies, organizations, etc.)
            for ent in doc.ents:
                if ent.label_ in ["ORG", "PRODUCT", "GPE"] and len(ent.text) > 2:
                    keywords_set.add(ent.text.lower())
            
            job_keywords_list = list(keywords_set)[:25]  # Limit to top 25 keywords
            logger.info(f"Auto-extracted keywords: {job_keywords_list[:10]}...")  # Log first 10
        else:
            job_keywords_list = [kw.strip() for kw in job_keywords.split(",")]

        # Score the resume
        result = score_resume(resume_text, job_keywords_list)
        
        # Add metadata to response
        result.update({
            "extraction_method": extraction_method,
            "processing_time": (datetime.now() - start_time).total_seconds(),
            "timestamp": datetime.now().isoformat(),
            "keywords_used": len(job_keywords_list),
            "text_length": len(resume_text)
        })
        
        logger.info(f"Scoring completed in {result['processing_time']:.2f}s, Score: {result.get('total_score', 'N/A')}")
        
        return jsonify(result)

    except Exception as e:
        logger.error(f"Error during scoring: {e}", exc_info=True)
        return jsonify({
            "error": "Internal server error during analysis. Please try again.",
            "processing_time": (datetime.now() - start_time).total_seconds()
        }), 500



    try:
        data = request.get_json()
        resume_b64 = data.get("resume")
        jd_b64 = data.get("jd")
        resume_name = data.get("resume_name", "resume.pdf")
        jd_name = data.get("jd_name", "jobdesc.pdf")

        if not resume_b64 or not jd_b64:
            return jsonify({"error": "Both resume and job description are required"}), 400

        # Decode base64 files and write to temp files
        resume_temp = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(resume_name)[-1])
        jd_temp = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(jd_name)[-1])

        resume_temp.write(base64.b64decode(resume_b64))
        jd_temp.write(base64.b64decode(jd_b64))
        resume_temp.close()
        jd_temp.close()

        # Extract text from both files
        resume_text, resume_method = extract_text_from_pdf(open(resume_temp.name, "rb")) \
            if resume_name.endswith(".pdf") else extract_text_from_docx(open(resume_temp.name, "rb"))
        
        jd_text, jd_method = extract_text_from_pdf(open(jd_temp.name, "rb")) \
            if jd_name.endswith(".pdf") else extract_text_from_docx(open(jd_temp.name, "rb"))

        if not resume_text.strip() or not jd_text.strip():
            return jsonify({"error": "Text extraction failed"}), 400

        # Extract keywords from JD
        doc = nlp(jd_text.lower())
        jd_keywords = set()
        for token in doc:
            if token.pos_ in ["NOUN", "PROPN"] and not token.is_stop and token.is_alpha and len(token.text) > 2:
                jd_keywords.add(token.text.strip())

        jd_keywords = list(jd_keywords)[:25]

        # Score
        result = score_resume(resume_text, jd_keywords)
        result.update({
            "extraction_methods": {
                "resume": resume_method,
                "jd": jd_method
            },
            "keywords_used": len(jd_keywords),
            "timestamp": datetime.now().isoformat()
        })

        return jsonify(result)
    
    except Exception as e:
        logger.error(f"Error in /match: {e}", exc_info=True)
        return jsonify({"error": "Server error during matching"}), 500



@app.route("/match", methods=["POST"])
def match_resume():
    try:
        if "resume" not in request.files or "job_description" not in request.form:
            return jsonify({"error": "Resume file and job description text are required"}), 400

        file = request.files["resume"]
        job_description = request.form["job_description"]

        if file.filename == "":
            return jsonify({"error": "No resume file selected"}), 400

        # Save temp file
        with tempfile.NamedTemporaryFile(delete=False) as tmp:
            file.save(tmp.name)
            file_path = tmp.name

        # Extract resume text
        if file.filename.lower().endswith(".pdf"):
            resume_text, _ = extract_text_from_pdf(open(file_path, "rb"))
        elif file.filename.lower().endswith(".docx"):
            resume_text, _ = extract_text_from_docx(open(file_path, "rb"))
        else:
            os.remove(file_path)
            return jsonify({"error": "Unsupported file type"}), 400

        os.remove(file_path)

        # ✅ Extract keywords from JD using extractor module
        job_keywords = extract_keywords(job_description, top_n=25)

        # Score resume
        result = score_resume(resume_text, job_keywords)
        result.update({
            "keywords_used": len(job_keywords),
            "timestamp": datetime.now().isoformat()
        })

        return jsonify(result)

    except Exception as e:
        logger.error(f"Error in /match route: {e}", exc_info=True)
        return jsonify({"error": "Server error during matching"}), 500
if __name__ == "__main__":
    app.run(debug=True, port=8000)
